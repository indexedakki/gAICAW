from docx import Document
import pandas as pd
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def replace_placeholder_in_word_from_excel(excel_file, doc, placeholder, modified_word_file_path):
    # Step 1: Read the Excel data
    df = pd.read_excel(excel_file)

    try:
        # Replace "Unnamed: X" column names with an empty string
        df.columns = ['' if 'Unnamed' in col else col for col in df.columns]
    except:
        pass

    # Step 3: Add the table at the end of the document
    table = doc.add_table(rows=1, cols=len(df.columns))
    for i, col_name in enumerate(df.columns):
        table.cell(0, i).text = str(col_name)
    for index, row in df.iterrows():
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)

    if 'distribution' in excel_file:
        # Access the XML element of the table
        table_element = table._tbl
        # Access the first row of the table
        first_row = table.rows[0]
        # Remove the first row from the table's XML element
        table_element.remove(first_row._tr)

        # Replace 0 with ''
        table.cell(0, 0).text = ''
        table.cell(0, 2).text = ''
        table.cell(0, 4).text = ''
        table.cell(0, 5).text = ''

        # Merge the first two cells in the first row
        table.cell(0, 1).merge(table.cell(0, 2))
        table.cell(0, 3).merge(table.cell(0, 4))
        

    # Function to add borders to a table
    def add_borders(table):
        for row in table.rows:
            for cell in row.cells:
                cell._tc.get_or_add_tcPr().append(parse_xml(r'<w:tcBorders %s><w:top w:val="single"/><w:left w:val="single"/><w:bottom w:val="single"/><w:right w:val="single"/></w:tcBorders>' % nsdecls('w')))

    # Add borders to the table
    add_borders(table)

    # Function to set font size and name for a table
    def set_font(table, font_name, font_size):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = font_name
                        run.font.size = Pt(font_size)

    # Set font size and name for the table
    set_font(table, "Arial", 10)

    # Function to make the text in a row bold
    def make_row_bold(row):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

    # Make the first row bold
    make_row_bold(table.rows[0])

    # Step 4: Find the placeholder and move the table
    placeholder = placeholder
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Move the table to the location of the placeholder
            table_el = table._element
            parent_el = paragraph._element.getparent()
            parent_el.insert(parent_el.index(paragraph._element) + 1, table_el)
            # Remove the placeholder paragraph
            paragraph._element.getparent().remove(paragraph._element)
            break

    # Save the modified Word document
    doc.save(modified_word_file_path)

def run_for_different_tables():
    word_file_path = 'C:\\Users\\2113196\\Downloads\\Detailed PSUR template_ICH E2C R2.docx'
    modified_word_file_path = 'C:\\Users\\2113196\\Downloads\\modified_word_file1.docx'
 
    doc = Document(word_file_path)

    excel_file = 'C:\\Users\\2113196\\Downloads\\clinical_input_table.xlsx'
    placeholder = '<@table 5.1@>'
    replace_placeholder_in_word_from_excel(excel_file, doc, placeholder, modified_word_file_path)

    placeholder = '<@table 5.8@>'
    excel_file = 'C:\\Users\\2113196\\Downloads\\Losartan_exposure_table_Formulation_PTY_50mg.xlsx'
    replace_placeholder_in_word_from_excel(excel_file, doc, placeholder, modified_word_file_path)

    placeholder = '<@table 5.9@>'
    excel_file = 'C:\\Users\\2113196\\Downloads\\Losartan_exposure_table_Region_PTY_50mg.xlsx'
    replace_placeholder_in_word_from_excel(excel_file, doc, placeholder, modified_word_file_path)

    placeholder = '<@table 5.10@>'
    excel_file = 'C:\\Users\\2113196\\Downloads\\LL_special_population.xlsx'
    replace_placeholder_in_word_from_excel(excel_file, doc, placeholder, modified_word_file_path)

    placeholder = '<@table 5.11@>'
    excel_file = 'C:\\Users\\2113196\\Downloads\\pediatric_case_distribution_reporting.xlsx'
    replace_placeholder_in_word_from_excel(excel_file, doc, placeholder, modified_word_file_path)

    placeholder = '<@table 5.12@>'
    excel_file = 'C:\\Users\\2113196\\Downloads\\geriatric_case_distribution_reporting.xlsx'
    replace_placeholder_in_word_from_excel(excel_file, doc, placeholder, modified_word_file_path)

    placeholder = '<@table 5.13@>'
    excel_file = 'C:\\Users\\2113196\\Downloads\\drugOverdose_case_distribution_reporting.xlsx'
    replace_placeholder_in_word_from_excel(excel_file, doc, placeholder, modified_word_file_path)

    placeholder = '<@table 5.14@>'
    excel_file = 'C:\\Users\\2113196\\Downloads\\drugMisuse_case_distribution_cumulative.xlsx'
    replace_placeholder_in_word_from_excel(excel_file, doc, placeholder, modified_word_file_path)

    placeholder = '<@table 5.15@>'
    excel_file = 'C:\\Users\\2113196\\Downloads\\offLabel_case_distribution_cumulative.xlsx'
    replace_placeholder_in_word_from_excel(excel_file, doc, placeholder, modified_word_file_path)


run_for_different_tables()

