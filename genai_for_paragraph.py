from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.parser import OxmlElement
from docx.oxml.parser import parse_xml


def insert_paragraph_after(paragraph, text=None, style=None):
    """Insert a new paragraph after the given paragraph."""
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style is not None:
        new_para.style = style
    return new_para

def replace_placeholder_with_paragraph(doc, placeholder, replacing_para):

    # Iterate through paragraphs to find the placeholder
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            insert_paragraph_after(paragraph, replacing_para)

            # Remove the placeholder paragraph
            paragraph._element.getparent().remove(paragraph._element)
            break
    

def run_for_different_paras():
    # Load the Word document
    doc = Document("C:\\Users\\2113196\\Downloads\\modified_word_file1.docx")

    placeholder = '<@para 5.11@>'
    replacing_para = open("C:\\Users\\2113196\\Downloads\\pediatric_case_seriousness_reporting", "r").read()
    replace_placeholder_with_paragraph(doc, placeholder, replacing_para)

    placeholder = '<@para 5.12@>'
    replacing_para = open("C:\\Users\\2113196\\Downloads\\geriatric_case_seriousness_reporting", "r").read()
    replace_placeholder_with_paragraph(doc, placeholder, replacing_para)

    placeholder = '<@para 5.13@>'
    replacing_para = open("C:\\Users\\2113196\\Downloads\\drugOverdose_case_seriousness_reporting", "r").read()
    replace_placeholder_with_paragraph(doc, placeholder, replacing_para)

    placeholder = '<@para 5.14@>'
    replacing_para = open("C:\\Users\\2113196\\Downloads\\drugMisuse_case_seriousness_cumulative", "r").read()
    replace_placeholder_with_paragraph(doc, placeholder, replacing_para)

    placeholder = '<@para 5.15@>'
    replacing_para = open("C:\\Users\\2113196\\Downloads\\offLabel_case_seriousness_cumulative", "r").read()
    replace_placeholder_with_paragraph(doc, placeholder, replacing_para)

    # Save the modified document
    doc.save("C:\\Users\\2113196\\Downloads\\result.docx")

run_for_different_paras()

