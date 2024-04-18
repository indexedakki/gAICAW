from docx import Document
import pandas as pd

# Step 1: Read the Excel data
excel_file = 'intestinal_performation_case_line_listing_table_cumulative.xlsx'
df = pd.read_excel(excel_file)

# Step 2: Create a Word document object
word_file = 'Detailed PSUR template_ICH E2C R2.docx'
doc = Document(word_file)

# Step 3: Add the table at the end of the document
table = doc.add_table(rows=1, cols=len(df.columns))
for i, col_name in enumerate(df.columns):
    table.cell(0, i).text = str(col_name)
for index, row in df.iterrows():
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = str(value)

# Step 4: Find the placeholder and move the table
placeholder = '<@table 5.2@>'
for paragraph in doc.paragraphs:
    if placeholder in paragraph.text:
        # Move the table to the location of the placeholder
        table_el = table._element
        parent_el = paragraph._element.getparent()
        parent_el.insert(parent_el.index(paragraph._element) + 1, table_el)
        # Remove the placeholder paragraph
        paragraph._element.getparent().remove(paragraph._element)
        break

# Save the modified Word document
doc.save('modified_word_file.docx')