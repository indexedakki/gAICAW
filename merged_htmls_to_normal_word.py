import re
import zipfile
import os
import tempfile
import shutil
from lxml import etree
import xml.etree.ElementTree as ET
from docx import Document

def extract_xml_from_docx(docx_path, folder):
    tmp_dir = f"C:\\Users\\2113196\\Downloads\\{folder}"
    with zipfile.ZipFile(docx_path, 'r') as zip_ref:
        zip_ref.extractall(tmp_dir)
    xml_path = os.path.join(tmp_dir, f'word\\document.xml')
    with open(xml_path, 'r', encoding='utf-8') as file:
        xml_content = file.read()
    return xml_content, tmp_dir

def split_document_section_wise(start_tag, end_tag):
    start_part = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
    <w:document
        xmlns:wpc="http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas"
        xmlns:cx="http://schemas.microsoft.com/office/drawing/2014/chartex"
        xmlns:cx1="http://schemas.microsoft.com/office/drawing/2015/9/8/chartex"
        xmlns:cx2="http://schemas.microsoft.com/office/drawing/2015/10/21/chartex"
        xmlns:cx3="http://schemas.microsoft.com/office/drawing/2016/5/9/chartex"
        xmlns:cx4="http://schemas.microsoft.com/office/drawing/2016/5/10/chartex"
        xmlns:cx5="http://schemas.microsoft.com/office/drawing/2016/5/11/chartex"
        xmlns:cx6="http://schemas.microsoft.com/office/drawing/2016/5/12/chartex"
        xmlns:cx7="http://schemas.microsoft.com/office/drawing/2016/5/13/chartex"
        xmlns:cx8="http://schemas.microsoft.com/office/drawing/2016/5/14/chartex"
        xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"
        xmlns:aink="http://schemas.microsoft.com/office/drawing/2016/ink"
        xmlns:am3d="http://schemas.microsoft.com/office/drawing/2017/model3d"
        xmlns:o="urn:schemas-microsoft-com:office:office"
        xmlns:oel="http://schemas.microsoft.com/office/2019/extlst"
        xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"
        xmlns:v="urn:schemas-microsoft-com:vml"
        xmlns:wp14="http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing"
        xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        xmlns:w10="urn:schemas-microsoft-com:office:word"
        xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        xmlns:w14="http://schemas.microsoft.com/office/word/2010/wordml"
        xmlns:w15="http://schemas.microsoft.com/office/word/2012/wordml"
        xmlns:w16cex="http://schemas.microsoft.com/office/word/2018/wordml/cex"
        xmlns:w16cid="http://schemas.microsoft.com/office/word/2016/wordml/cid"
        xmlns:w16="http://schemas.microsoft.com/office/word/2018/wordml"
        xmlns:w16sdtdh="http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash"
        xmlns:w16se="http://schemas.microsoft.com/office/word/2015/wordml/symex"
        xmlns:wpg="http://schemas.microsoft.com/office/word/2010/wordprocessingGroup"
        xmlns:wpi="http://schemas.microsoft.com/office/word/2010/wordprocessingInk"
        xmlns:wne="http://schemas.microsoft.com/office/word/2006/wordml"
        xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape" mc:Ignorable="w14 w15 w16se w16cid w16 w16cex w16sdtdh wp14">
        <w:body>
            <w:p w14:paraId="4F127794" w14:textId="77777777" w:rsidR="009710B7" w:rsidRDefault="009710B7" w:rsidP="009710B7">
                
                
                <w:r w:rsidRPr="00B41E60">
                    """
    end_part = """
            </w:p>
            
        </w:body>
    </w:document>"""

    def extract_text(input_string, start_tag, end_tag):
        occurrence = 1
        start_tag = start_tag.strip()
        end_tag = end_tag.strip()
        

        print("occurence-",occurrence)
        print("start_tag -", start_tag, "end_tag-", end_tag )
        splits = input_string.split(start_tag)
        print("len", len(splits))

        if len(splits) <= occurrence:
            print("Not enough instances found")

        rest_of_string = start_tag.join(splits[occurrence:])
        end_index = rest_of_string.find(end_tag)

        if end_index == -1:
            print(f"No {end_tag} found after the {occurrence}th occurrence of {start_tag}")

        return rest_of_string[:end_index + len(end_tag)]

    with open("C:\\Users\\2113196\\Downloads\\full_psur\\word\\document.xml", 'r', encoding='utf-8') as f:
        xml_content = f.read()
    text = xml_content

    # To see extracted text
    with open('C:\\Users\\2113196\\Downloads\\full_document_xml.txt', 'w', encoding='utf-8') as file:
        file.write(text)
    
    extracted_text = extract_text(text, start_tag, end_tag)

    with open('C:\\Users\\2113196\\Downloads\\extracted_xml.txt', 'w', encoding='utf-8') as file:
        file.write(extracted_text)

    remove_extra_end_part = "</w:r>"
    index = extracted_text.rfind(remove_extra_end_part)
    modified_text = extracted_text[:index + len(remove_extra_end_part)]

    middle_text = "<w:t>" + " "
    final_text = start_part + middle_text + modified_text + end_part

    with open('C:\\Users\\2113196\\Downloads\\extracted_text.txt', 'w', encoding='utf-8') as file:
        file.write(final_text)
        
    # Write the modified data back to the file
    with open("C:\\Users\\2113196\\Downloads\\section_wise_temporary\\word\\document.xml", 'w', encoding='utf-8') as f:
        f.write(final_text)


def xml_to_docx(xml_content, output_path):
    tmp_dir = "C:\\Users\\2113196\\Downloads\\template"
    xml_path = os.path.join(tmp_dir, 'word\\document.xml')
    with open(xml_path, 'w', encoding='utf-8') as file:
        file.write(xml_content)
    with zipfile.ZipFile(output_path, 'w') as zip_ref:
        for root, dirs, files in os.walk(tmp_dir):
            for file in files:
                zip_ref.write(os.path.join(root, file),
                              os.path.relpath(os.path.join(root, file), tmp_dir))
    
def create_template_folders_with_xml():
    folder = "full_psur"
    extract_xml_from_docx("C:\\Users\\2113196\\Downloads\\merged.docx", folder)
    folder = "section_wise_temporary"
    extract_xml_from_docx("C:\\Users\\2113196\\Downloads\\merged.docx", folder)


def splitting_and_converting_to_different_word():   
    def run_xml_to_docs(section_name):
        tmp_dir = "C:\\Users\\2113196\\Downloads\\section_wise_temporary"
        xml_path = os.path.join(tmp_dir, 'word\\document.xml')
        with open(xml_path, 'r', encoding='utf-8') as file:
            xml_content = file.read()

        root = etree.fromstring(xml_content.encode('utf-8'))

        xml_to_docx(etree.tostring(root, pretty_print=True).decode('utf-8'), f"C:\\Users\\2113196\\Downloads\\{section_name}.docx")

    def run_for_different_sections(from_section , to_section, section_name):
        start_tag = from_section
        end_tag = to_section
        split_document_section_wise(start_tag, end_tag)
        run_xml_to_docs(section_name)

    # def extract_headers(doc_path):
    #     document = Document(doc_path)
    #     headers = []
    #     for paragraph in document.paragraphs:
    #         if paragraph.style.name.startswith('Heading 1'):
    #             headers.append(paragraph.text)
    #     return headers
    
    # # Extracting section headers from document
    # doc_path = "C:\\Users\\2113196\\Downloads\\result.docx"
    # headers = extract_headers(doc_path)
    # print(headers)

    start_tag = "word_document_start"
    end_tag = "word_document_end"
    run_for_different_sections(start_tag, end_tag, f"whole_doc")

def adding_borders_to_all_tables():
    from docx import Document
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml

    # Load the document
    doc = Document("C:\\Users\\2113196\\Downloads\\whole_doc.docx")

    # Define the border style using XML
    border_xml = '''
    <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
    <w:top w:val="single"/>
    <w:left w:val="single"/>
    <w:bottom w:val="single"/>
    <w:right w:val="single"/>
    </w:tcBorders>
    '''

    # Apply the border to each table and each cell
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell._tc.get_or_add_tcPr().append(parse_xml(border_xml))

    # Save the document
    doc.save("C:\\Users\\2113196\\Downloads\\whole_doc1.docx")

create_template_folders_with_xml()        
splitting_and_converting_to_different_word()
adding_borders_to_all_tables()
